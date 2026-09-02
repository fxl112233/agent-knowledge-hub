from __future__ import annotations

from pathlib import Path

import pytest

from agents.doc_parser_agent import DocParserAgent, DocumentParseError
from tests.conftest import FakeChatModel


@pytest.mark.asyncio
async def test_text_encoding_and_stable_chunk_ids(tmp_path: Path) -> None:
    path = tmp_path / "sample.txt"
    content = "张三负责 Atlas 项目。\n" * 700
    path.write_bytes(content.encode("gb18030"))
    parser = DocParserAgent()
    first = await parser.parse(str(path), doc_id="doc-1", version=1)
    second = await parser.parse(str(path), doc_id="doc-1", version=2)
    assert len(first) > 1
    assert [chunk.chunk_id for chunk in first] == [chunk.chunk_id for chunk in second]
    assert all(chunk.metadata["version"] in {1, 2} for chunk in first + second)
    assert all(chunk.metadata["chunk_level"] == "child" for chunk in first)
    assert all(chunk.metadata["parent_id"] for chunk in first)
    assert all(chunk.metadata["parent_content"] for chunk in first)
    assert [chunk.metadata["parent_id"] for chunk in first] == [
        chunk.metadata["parent_id"] for chunk in second
    ]
    assert all(
        chunk.metadata["parent_token_start"] <= chunk.metadata["token_start"]
        < chunk.metadata["token_end"] <= chunk.metadata["parent_token_end"]
        or chunk.metadata["token_start"] < chunk.metadata["parent_token_end"]
        for chunk in first
    )


@pytest.mark.asyncio
async def test_markdown_section_ids_survive_unrelated_insert(tmp_path: Path) -> None:
    path = tmp_path / "sections.md"
    path.write_text("## Alpha\nalpha fact\n\n## Beta\nbeta fact\n", encoding="utf-8")
    parser = DocParserAgent()
    first = await parser.parse(str(path), doc_id="markdown", version=1)
    first_by_section = {chunk.metadata["section"]: chunk.chunk_id for chunk in first}

    path.write_text(
        "## Inserted\nnew fact\n\n## Alpha\nalpha fact\n\n## Beta\nbeta fact\n",
        encoding="utf-8",
    )
    second = await parser.parse(str(path), doc_id="markdown", version=2)
    second_by_section = {chunk.metadata["section"]: chunk.chunk_id for chunk in second}
    assert second_by_section["Alpha"] == first_by_section["Alpha"]
    assert second_by_section["Beta"] == first_by_section["Beta"]


@pytest.mark.asyncio
async def test_csv_and_xlsx_keep_coordinates(tmp_path: Path) -> None:
    csv_path = tmp_path / "table.csv"
    csv_path.write_text("name,value\nrevenue,123\n", encoding="utf-8")
    csv_chunks = await DocParserAgent().parse(str(csv_path), doc_id="csv")
    assert csv_chunks[0].metadata["row_start"] == 1
    assert "revenue" in csv_chunks[0].content

    from openpyxl import Workbook

    xlsx_path = tmp_path / "table.xlsx"
    workbook = Workbook()
    workbook.active.append(["name", "value"])
    workbook.active.append(["growth", 12])
    workbook.save(xlsx_path)
    xlsx_chunks = await DocParserAgent().parse(str(xlsx_path), doc_id="xlsx")
    assert xlsx_chunks[0].metadata["sheet"] == "Sheet"
    assert xlsx_chunks[0].metadata["row_start"] == 1


@pytest.mark.asyncio
async def test_rejects_missing_unsupported_and_empty_files(tmp_path: Path) -> None:
    parser = DocParserAgent()
    with pytest.raises(DocumentParseError):
        await parser.parse(str(tmp_path / "missing.txt"))
    unsupported = tmp_path / "x.exe"
    unsupported.write_bytes(b"x")
    with pytest.raises(DocumentParseError):
        await parser.parse(str(unsupported))
    empty = tmp_path / "empty.md"
    empty.write_text("", encoding="utf-8")
    with pytest.raises(DocumentParseError):
        await parser.parse(str(empty))


@pytest.mark.asyncio
async def test_pdf_docx_image_and_batch_parsing(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    from docx import Document
    from docx.shared import Inches
    from PIL import Image
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen.canvas import Canvas

    pdf_path = tmp_path / "report.pdf"
    canvas = Canvas(str(pdf_path), pagesize=A4)
    canvas.drawString(
        40, A4[1] - 50, "Atlas annual report contains enough searchable digital text for parsing."
    )
    canvas.save()

    image_path = tmp_path / "chart.png"
    Image.new("RGB", (240, 100), "white").save(image_path)
    docx_path = tmp_path / "report.docx"
    document = Document()
    document.add_heading("Atlas", level=1)
    document.add_paragraph("The project owner is Alice.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "100"
    document.add_picture(str(image_path), width=Inches(1))
    document.save(docx_path)

    monkeypatch.setattr(DocParserAgent, "_ocr_image", staticmethod(lambda _image: "chart revenue 100"))
    parser = DocParserAgent()
    pdf = await parser.parse(str(pdf_path), doc_id="pdf")
    docx = await parser.parse(str(docx_path), doc_id="docx")
    image = await parser.parse(str(image_path), doc_id="image")
    batch = await parser.parse_batch([str(pdf_path), str(image_path)], doc_ids=["p", "i"])
    assert pdf[0].metadata["page"] == 1
    assert any(chunk.metadata.get("kind") == "table" for chunk in docx)
    assert any(chunk.metadata.get("kind") == "image" for chunk in docx)
    assert "chart revenue" in image[0].content
    assert {chunk.doc_id for chunk in batch} == {"p", "i"}
    with pytest.raises(ValueError):
        await parser.parse_batch([str(pdf_path)], doc_ids=[])


@pytest.mark.asyncio
async def test_image_vision_description_and_invalid_image(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    from PIL import Image

    image_path = tmp_path / "chart.jpg"
    Image.new("RGB", (100, 80), "white").save(image_path)
    monkeypatch.setattr(DocParserAgent, "_ocr_image", staticmethod(lambda _image: "OCR text"))
    monkeypatch.setattr("agents.doc_parser_agent.settings.llm_supports_vision", True)
    parser = DocParserAgent(vision_model=FakeChatModel(["visual description"]))
    chunks = await parser.parse(str(image_path), doc_id="vision")
    assert "OCR text" in chunks[0].content
    assert "visual description" in chunks[0].content

    invalid = tmp_path / "invalid.png"
    invalid.write_text("not an image", encoding="utf-8")
    with pytest.raises(DocumentParseError):
        await parser.parse(str(invalid))
