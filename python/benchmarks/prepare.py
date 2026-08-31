"""Deterministically convert public datasets into benchmark documents and JSONL cases."""

from __future__ import annotations

import csv
import json
import random
import re
import shutil
import textwrap
from collections import Counter, defaultdict
from collections.abc import Iterable
from pathlib import Path
from typing import Any

from benchmarks.data import PREPARED_ROOT, RAW_ROOT, sha256_file, verify_datasets

SEED = 42
MULTIHOP_DOCUMENT_LIMIT = 100


def _json_files(root: Path) -> Iterable[Path]:
    yield from sorted(path for path in root.rglob("*") if path.suffix.lower() in {".json", ".jsonl"})


def _load_records(path: Path) -> list[Any]:
    text = path.read_text(encoding="utf-8")
    try:
        if path.suffix.lower() == ".jsonl":
            return [json.loads(line) for line in text.splitlines() if line.strip()]
        value = json.loads(text)
        if isinstance(value, list):
            return value
        if isinstance(value, dict):
            for key in ("data", "records", "questions", "examples"):
                if isinstance(value.get(key), list):
                    return value[key]
            return [value]
    except json.JSONDecodeError:
        try:
            return [json.loads(line) for line in text.splitlines() if line.strip()]
        except json.JSONDecodeError:
            return []
    return []


def _write_jsonl(path: Path, records: Iterable[dict[str, Any]]) -> int:
    values = list(records)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        "".join(json.dumps(record, ensure_ascii=False, sort_keys=True) + "\n" for record in values),
        encoding="utf-8",
    )
    return len(values)


def _safe_identifier(value: object, fallback: str) -> str:
    identifier = re.sub(r"[^a-zA-Z0-9_.-]+", "-", str(value or "")).strip("-.")
    return identifier[:100] or fallback


def _stratified_sample(
    records: list[dict[str, Any]],
    size: int,
    category_key: str,
    *,
    seed: int = SEED,
) -> list[dict[str, Any]]:
    if len(records) <= size:
        return records.copy()
    groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for record in records:
        groups[str(record.get(category_key, "unknown"))].append(record)
    generator = random.Random(seed)
    for values in groups.values():
        generator.shuffle(values)
    selected: list[dict[str, Any]] = []
    names = sorted(groups)
    while len(selected) < size:
        progressed = False
        for name in names:
            if groups[name] and len(selected) < size:
                selected.append(groups[name].pop())
                progressed = True
        if not progressed:
            break
    return selected


def _select_evidence_documents(cases: list[dict[str, Any]], limit: int = MULTIHOP_DOCUMENT_LIMIT) -> set[str]:
    """Select a deterministic, high-coverage document subset for retrieval evaluation."""
    evidence_frequency = Counter(str(doc_id) for case in cases for doc_id in case.get("evidence_doc_ids", []))
    return {
        doc_id
        for doc_id, _count in sorted(evidence_frequency.items(), key=lambda item: (-item[1], item[0]))[:limit]
    }


def prepare_multihop() -> dict[str, Any]:
    raw_root = RAW_ROOT / "multihop_rag"
    candidates = [(path, _load_records(path)) for path in _json_files(raw_root)]
    corpus_path, corpus = max(
        (
            (path, rows)
            for path, rows in candidates
            if rows and isinstance(rows[0], dict) and ({"body", "title"} <= set(rows[0]) or "text" in rows[0])
        ),
        key=lambda item: len(item[1]),
    )
    qa_path, questions = max(
        (
            (path, rows)
            for path, rows in candidates
            if rows
            and isinstance(rows[0], dict)
            and ("query" in rows[0] or "question" in rows[0])
            and "answer" in rows[0]
        ),
        key=lambda item: len(item[1]),
    )
    output = PREPARED_ROOT / "multihop_rag"
    documents = output / "documents"
    documents.mkdir(parents=True, exist_ok=True)
    title_to_id: dict[str, str] = {}
    url_to_id: dict[str, str] = {}
    article_records: list[dict[str, Any]] = []
    for index, article in enumerate(corpus):
        title = str(article.get("title") or article.get("name") or f"article-{index:04d}")
        doc_id = _safe_identifier(article.get("id") or article.get("url") or title, f"article-{index:04d}")
        if doc_id in title_to_id.values():
            doc_id = f"{doc_id}-{index:04d}"
        body = str(article.get("body") or article.get("text") or article.get("content") or "")
        source = str(article.get("source") or article.get("url") or "")
        url = str(article.get("url") or "")
        title_to_id[title] = doc_id
        if url:
            url_to_id[url] = doc_id
        article_records.append(
            {
                "doc_id": doc_id,
                "title": title,
                "body": body,
                "source": source,
                "path": f"documents/{doc_id}.md",
            }
        )

    all_cases = []
    for index, item in enumerate(questions):
        evidence = item.get("evidence_list") or item.get("evidence") or item.get("supporting_facts") or []
        evidence_titles: list[str] = []
        evidence_doc_ids: set[str] = set()
        if isinstance(evidence, list):
            for value in evidence:
                if isinstance(value, str):
                    evidence_titles.append(value)
                    if value in title_to_id:
                        evidence_doc_ids.add(title_to_id[value])
                elif isinstance(value, dict):
                    title = str(value.get("title") or value.get("source") or value.get("document") or "")
                    url = str(value.get("url") or "")
                    evidence_titles.append(title)
                    doc_id = url_to_id.get(url) or title_to_id.get(title)
                    if doc_id:
                        evidence_doc_ids.add(doc_id)
                elif isinstance(value, list) and value:
                    title = str(value[0])
                    evidence_titles.append(title)
                    if title in title_to_id:
                        evidence_doc_ids.add(title_to_id[title])
        category = str(item.get("question_type") or item.get("type") or "unknown")
        all_cases.append(
            {
                "sample_id": str(item.get("id") or item.get("qid") or f"multihop-{index:04d}"),
                "question": str(item.get("query") or item.get("question") or ""),
                "answer": item.get("answer", ""),
                "category": category,
                "evidence_titles": sorted(set(filter(None, evidence_titles))),
                "evidence_doc_ids": sorted(evidence_doc_ids),
                "unanswerable": category == "null_query" or not evidence_doc_ids,
            }
        )

    selected_ids = _select_evidence_documents(all_cases)
    selected_articles = [record for record in article_records if record["doc_id"] in selected_ids]
    cases = [
        case
        for case in all_cases
        if case["evidence_doc_ids"] and set(case["evidence_doc_ids"]) <= selected_ids
    ]
    null_cases = [case for case in all_cases if case["unanswerable"]]
    answer_sample = _stratified_sample(cases + null_cases, 300, "category")

    expected_paths = {documents / f"{record['doc_id']}.md" for record in selected_articles}
    for stale_path in documents.glob("*.md"):
        if stale_path not in expected_paths:
            stale_path.unlink()
    document_records = []
    for record in selected_articles:
        path = documents / f"{record['doc_id']}.md"
        path.write_text(
            f"# {record['title']}\n\n{record['body']}\n\nSource: {record['source']}\n",
            encoding="utf-8",
        )
        document_records.append(
            {
                "doc_id": record["doc_id"],
                "title": record["title"],
                "path": path.relative_to(output).as_posix(),
                "sha256": sha256_file(path),
            }
        )

    _write_jsonl(output / "documents.jsonl", document_records)
    _write_jsonl(output / "retrieval.jsonl", cases)
    _write_jsonl(output / "answer-300.jsonl", answer_sample)
    (output / "selection.json").write_text(
        json.dumps(
            {
                "schema_version": 1,
                "seed": SEED,
                "selection_method": "top evidence-document frequency; doc_id ascending tie-break",
                "source_document_count": len(article_records),
                "document_limit": MULTIHOP_DOCUMENT_LIMIT,
                "selected_document_ids": sorted(selected_ids),
                "retrieval_sample_ids": [case["sample_id"] for case in cases],
                "answer_sample_ids": [case["sample_id"] for case in answer_sample],
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    return {
        "documents": len(document_records),
        "source_documents": len(article_records),
        "retrieval_cases": len(cases),
        "answer_cases": len(answer_sample),
        "unanswerable_answer_cases": sum(bool(case["unanswerable"]) for case in answer_sample),
        "source_files": 2,
        "corpus_source": corpus_path.relative_to(raw_root).as_posix(),
        "qa_source": qa_path.relative_to(raw_root).as_posix(),
    }


def _tat_context(document: dict[str, Any]) -> tuple[list[list[str]], list[str]]:
    table_value = document.get("table", [])
    if isinstance(table_value, dict):
        table_value = table_value.get("table") or table_value.get("data") or []
    table = [[str(cell or "") for cell in row] for row in table_value if isinstance(row, list)]
    paragraphs = document.get("paragraphs") or document.get("paragraph") or []
    texts = []
    for paragraph in paragraphs:
        texts.append(str(paragraph.get("text", "")) if isinstance(paragraph, dict) else str(paragraph))
    return table, texts


def _tat_questions(document: dict[str, Any]) -> list[dict[str, Any]]:
    value = document.get("questions") or document.get("qa") or []
    return [item for item in value if isinstance(item, dict)]


def _write_tat_document(path: Path, table: list[list[str]], paragraphs: list[str], format_name: str) -> None:
    if format_name == "xlsx":
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Table"
        for row in table:
            sheet.append(row)
        text_sheet = workbook.create_sheet("Paragraphs")
        for index, paragraph in enumerate(paragraphs, start=1):
            text_sheet.append([index, paragraph])
        workbook.save(path)
    elif format_name == "docx":
        from docx import Document

        document = Document()
        document.add_heading("Financial context", level=1)
        if table:
            word_table = document.add_table(rows=len(table), cols=max(map(len, table)))
            for row_index, row in enumerate(table):
                for column_index, value in enumerate(row):
                    word_table.cell(row_index, column_index).text = value
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        document.save(path)
    else:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen.canvas import Canvas

        canvas = Canvas(str(path), pagesize=A4)
        width, height = A4
        y = height - 40
        lines = [" | ".join(row) for row in table] + paragraphs
        for line in lines:
            for start in range(0, len(line), 100):
                canvas.drawString(35, y, line[start : start + 100])
                y -= 14
                if y < 40:
                    canvas.showPage()
                    y = height - 40
        canvas.save()


def prepare_tatqa() -> dict[str, Any]:
    raw_root = RAW_ROOT / "tat_qa"
    dev_candidates = []
    for path in _json_files(raw_root):
        if "dev" in path.name.lower():
            rows = _load_records(path)
            if rows and isinstance(rows[0], dict) and _tat_questions(rows[0]):
                dev_candidates.append((path, rows))
    source_path, documents = max(dev_candidates, key=lambda item: len(item[1]))
    flattened = []
    for document_index, document in enumerate(documents):
        table_uid = document.get("table", {}).get("uid") if isinstance(document.get("table"), dict) else None
        doc_uid = str(document.get("uid") or table_uid or f"tat-doc-{document_index:04d}")
        for question in _tat_questions(document):
            flattened.append(
                {
                    "document": document,
                    "doc_uid": doc_uid,
                    "question": question,
                    "category": str(question.get("answer_type") or question.get("type") or "unknown"),
                }
            )
    sample = _stratified_sample(flattened, 300, "category")
    output = PREPARED_ROOT / "tat_qa"
    document_root = output / "documents"
    document_root.mkdir(parents=True, exist_ok=True)
    created: dict[str, dict[str, Any]] = {}
    cases = []
    formats = ("xlsx", "pdf", "docx")
    for index, row in enumerate(sample):
        doc_uid = row["doc_uid"]
        if doc_uid not in created:
            table, paragraphs = _tat_context(row["document"])
            format_name = formats[len(created) % len(formats)]
            path = document_root / f"{_safe_identifier(doc_uid, f'tat-{index:04d}')}.{format_name}"
            _write_tat_document(path, table, paragraphs, format_name)
            created[doc_uid] = {
                "doc_id": doc_uid,
                "path": path.relative_to(output).as_posix(),
                "format": format_name,
                "sha256": sha256_file(path),
            }
        question = row["question"]
        cases.append(
            {
                "sample_id": str(question.get("uid") or question.get("id") or f"tat-{index:04d}"),
                "doc_id": doc_uid,
                "question": str(question.get("question") or question.get("query") or ""),
                "answer": _canonical_answer(question.get("answer", "")),
                "derivation": str(question.get("derivation") or ""),
                "answer_type": row["category"],
                "scale": str(question.get("scale") or ""),
            }
        )
    _write_jsonl(output / "documents.jsonl", created.values())
    _write_jsonl(output / "answer-300.jsonl", cases)
    return {
        "documents": len(created),
        "answer_cases": len(cases),
        "source": source_path.relative_to(raw_root).as_posix(),
    }


def _passage_text(value: Any) -> str:
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        return str(value.get("text") or value.get("content") or value.get("passage") or value)
    return str(value)


def _canonical_answer(value: Any) -> str:
    if isinstance(value, list):
        return " ".join(str(item) for item in value)
    return str(value or "")


def prepare_rgb() -> dict[str, Any]:
    raw_root = RAW_ROOT / "rgb" / "data"
    source_groups = {
        "noise_40": _load_records(raw_root / "zh_refine.json"),
        "rejection": _load_records(raw_root / "zh_refine.json"),
        "integration": _load_records(raw_root / "zh_int.json"),
        "counterfactual": _load_records(raw_root / "zh_fact.json"),
    }
    sample: list[dict[str, Any]] = []
    generator = random.Random(SEED)
    for category, values in source_groups.items():
        indices = list(range(len(values)))
        generator.shuffle(indices)
        for source_index in indices[:50]:
            item = values[source_index]
            positive = list(item.get("positive") or [])
            negative = list(item.get("negative") or [])
            if category == "noise_40":
                passages = [*positive[:3], *negative[:2]]
            elif category == "rejection":
                passages = negative[:5]
            elif category == "integration":
                passages = [*positive, *negative[: max(1, round(len(positive) * 2 / 3))]]
            else:
                passages = [*(item.get("positive_wrong") or []), *positive, *negative[:2]]
            sample.append(
                {
                    "item": item,
                    "source_index": source_index,
                    "category": category,
                    "question": str(item.get("query") or item.get("question") or ""),
                    "answer": _canonical_answer(item.get("answer")),
                    "context": "\n\n".join(_passage_text(passage) for passage in passages),
                    "unanswerable": category == "rejection",
                }
            )
    output = PREPARED_ROOT / "rgb"
    document_root = output / "documents"
    document_root.mkdir(parents=True, exist_ok=True)
    cases = []
    documents = []
    for index, row in enumerate(sample):
        source_id = _safe_identifier(row["item"].get("id") or row["item"].get("qid"), f"rgb-{index:04d}")
        sample_id = f"{row['category']}-{source_id}"
        doc_id = f"rgb-doc-{sample_id}"
        path = document_root / f"{doc_id}.md"
        path.write_text(f"# 中文参考资料\n\n{row['context']}\n", encoding="utf-8")
        documents.append(
            {"doc_id": doc_id, "path": path.relative_to(output).as_posix(), "sha256": sha256_file(path)}
        )
        cases.append(
            {
                "sample_id": sample_id,
                "doc_id": doc_id,
                "question": row["question"],
                "answer": row["answer"],
                "category": row["category"],
                "source": f"zh_{row['category']}.json",
                "unanswerable": row["unanswerable"],
            }
        )
    _write_jsonl(output / "documents.jsonl", documents)
    _write_jsonl(output / "answer-200.jsonl", cases)
    return {"documents": len(documents), "answer_cases": len(cases)}


def _sample_text(index: int, cases: list[dict[str, Any]]) -> tuple[str, list[list[str]], str]:
    case = cases[index % len(cases)]
    sample_id = str(case["sample_id"])
    answer = str(case.get("answer", ""))
    answer_type = str(case.get("answer_type", "unknown"))
    text = f"TAT-QA sample {sample_id}. Question: {case['question']} Reference answer: {answer}."
    table = [
        ["Field", "Value", "Source"],
        ["Reference Answer", answer, "TAT-QA"],
        ["Answer Type", answer_type, sample_id],
    ]
    return text, table, sample_id


def _generate_parse70(output: Path) -> dict[str, int]:
    from docx import Document
    from openpyxl import Workbook
    from PIL import Image, ImageDraw, ImageFilter
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen.canvas import Canvas

    document_root = output / "documents"
    if document_root.exists():
        shutil.rmtree(document_root)
    document_root.mkdir(parents=True)
    references = []
    tat_path = PREPARED_ROOT / "tat_qa" / "answer-300.jsonl"
    tat_cases = [
        json.loads(line) for line in tat_path.read_text(encoding="utf-8").splitlines() if line.strip()
    ]
    if len(tat_cases) < 70:
        raise RuntimeError("AKH-Parse-70 requires at least 70 prepared TAT-QA cases")
    categories = ("pdf", "docx", "xlsx", "csv", "txt", "md", "image")
    for category_index, category in enumerate(categories):
        for offset in range(10):
            index = category_index * 10 + offset
            text, table, source_sample_id = _sample_text(index, tat_cases)
            expected_cells: list[str] = []
            condition = "digital"
            if category == "pdf":
                path = document_root / f"parse-{index:02d}.pdf"
                canvas = Canvas(str(path), pagesize=A4)
                text_lines = textwrap.wrap(text, width=92)
                for line_index, line in enumerate(text_lines):
                    canvas.drawString(50, A4[1] - 60 - line_index * 18, line)
                table_y = A4[1] - 100 - max(0, len(text_lines) - 1) * 18
                for row_index, row in enumerate(table):
                    canvas.drawString(50, table_y - row_index * 18, " | ".join(row))
                canvas.save()
                expected_cells = [cell for row in table for cell in row]
            elif category == "docx":
                path = document_root / f"parse-{index:02d}.docx"
                document = Document()
                document.add_heading("Atlas Report", level=1)
                document.add_paragraph(text)
                word_table = document.add_table(rows=len(table), cols=3)
                for row_index, row in enumerate(table):
                    for column_index, value in enumerate(row):
                        word_table.cell(row_index, column_index).text = value
                document.save(path)
                expected_cells = [cell for row in table for cell in row]
            elif category == "xlsx":
                path = document_root / f"parse-{index:02d}.xlsx"
                workbook = Workbook()
                sheet = workbook.active
                sheet.title = "Metrics"
                for row in table:
                    sheet.append(row)
                notes = workbook.create_sheet("Notes")
                notes["A1"] = text
                workbook.save(path)
                expected_cells = [cell for row in table for cell in row]
            elif category == "csv":
                path = document_root / f"parse-{index:02d}.csv"
                with path.open("w", encoding="utf-8", newline="") as handle:
                    writer = csv.writer(handle)
                    writer.writerows(table)
                expected_cells = [cell for row in table for cell in row]
                text = "\n".join(",".join(row) for row in table)
            elif category == "txt":
                path = document_root / f"parse-{index:02d}.txt"
                path.write_text(text, encoding="utf-8")
            elif category == "md":
                path = document_root / f"parse-{index:02d}.md"
                path.write_text(f"# Atlas Report\n\n{text}\n", encoding="utf-8")
            else:
                extension = "png" if offset < 5 else "jpg"
                path = document_root / f"parse-{index:02d}.{extension}"
                image = Image.new("RGB", (1800, 420), "white")
                drawer = ImageDraw.Draw(image)
                rendered = "\n".join(textwrap.wrap(text, width=105))
                drawer.multiline_text((45, 55), rendered, fill="black", font_size=28, spacing=12)
                condition = "clear_ocr"
                if offset >= 5:
                    image = image.rotate(2.5, expand=False, fillcolor="white").filter(
                        ImageFilter.GaussianBlur(0.7)
                    )
                    condition = "disturbed_ocr"
                image.save(path, quality=48 if extension == "jpg" else None)
            expected_text = text
            if expected_cells and category != "csv":
                expected_text += "\n" + "\n".join(" | ".join(row) for row in table)
            references.append(
                {
                    "sample_id": f"parse-{index:02d}",
                    "category": category,
                    "condition": condition,
                    "path": path.relative_to(output).as_posix(),
                    "expected_text": expected_text,
                    "expected_cells": expected_cells,
                    "source_dataset": "TAT-QA",
                    "source_sample_id": source_sample_id,
                    "sha256": sha256_file(path),
                }
            )
    _write_jsonl(output / "references.jsonl", references)
    return {"documents": len(references), "categories": len(categories)}


def _generate_cdc300(output: Path) -> dict[str, int]:
    snapshot_root = output / "snapshots"
    if snapshot_root.exists():
        shutil.rmtree(snapshot_root)
    snapshot_root.mkdir(parents=True)
    documents = []
    events = []
    for index in range(100):
        doc_id = f"cdc-{index:03d}"
        directory = snapshot_root / doc_id
        directory.mkdir()
        sections = [
            f"## Stable section {part:02d}\nDocument {index} stable fact {part}: "
            f"the recorded value is {index * 100 + part} and remains unchanged."
            for part in range(20)
        ]
        versions = [
            sections,
            [
                *sections[:2],
                f"## Stable section 02\nDocument {index} modified fact 2: "
                f"the recorded value is {index * 100 + 2}-v2.",
                *sections[3:],
            ],
            [
                *sections[:2],
                f"## Stable section 02\nDocument {index} modified fact 2: "
                f"the recorded value is {index * 100 + 2}-v2.",
                *sections[3:],
                f"## Added section 20\nDocument {index} added fact 20: "
                f"the recorded value is {index * 100 + 20}.",
            ],
            [
                *sections[:2],
                f"## Stable section 02\nDocument {index} modified fact 2: "
                f"the recorded value is {index * 100 + 2}-v2.",
                *sections[4:],
                f"## Added section 20\nDocument {index} added fact 20: "
                f"the recorded value is {index * 100 + 20}.",
            ],
        ]
        for version, values in enumerate(versions, start=1):
            path = directory / f"v{version}.md"
            path.write_text("\n\n".join(values) + "\n", encoding="utf-8")
        documents.append({"doc_id": doc_id, "initial": (directory / "v1.md").relative_to(output).as_posix()})
        changes = (("modified", 19), ("added", 20), ("deleted", 20))
        for sequence, (change, expected_unchanged) in enumerate(changes, start=1):
            target = directory / f"v{sequence + 1}.md"
            events.append(
                {
                    "event_id": f"cdc-event-v2-{index:03d}-{sequence}",
                    "doc_id": doc_id,
                    "operation": "UPDATE",
                    "change": change,
                    "sequence": sequence,
                    "expected_unchanged_chunks": expected_unchanged,
                    "snapshot": target.relative_to(output).as_posix(),
                    "sha256": sha256_file(target),
                }
            )
    _write_jsonl(output / "documents.jsonl", documents)
    _write_jsonl(output / "events-300.jsonl", events)
    return {"documents": len(documents), "events": len(events)}


def prepare_all(*, allow_unverified: bool = False) -> dict[str, Any]:
    verification = verify_datasets()
    if not allow_unverified and not all(verification.values()):
        failed = ", ".join(name for name, valid in verification.items() if not valid)
        raise RuntimeError(f"dataset integrity check failed: {failed}")
    PREPARED_ROOT.mkdir(parents=True, exist_ok=True)
    summary = {
        "seed": SEED,
        "verification": verification,
        "multihop_rag": prepare_multihop(),
        "tat_qa": prepare_tatqa(),
        "rgb": prepare_rgb(),
        "akh_parse_70": _generate_parse70(PREPARED_ROOT / "akh_parse_70"),
        "akh_cdc_300": _generate_cdc300(PREPARED_ROOT / "akh_cdc_300"),
    }
    (PREPARED_ROOT / "summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    return summary
